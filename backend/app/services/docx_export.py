"""Export a generated paper to an IEEE-formatted .docx (two-column, A4).

Layout mirrors the IEEE conference template:
  - full-width title + authors block,
  - a continuous section break into two columns for the Abstract, body, and
    References,
  - Times New Roman, 10pt body, justified, roman-numeral section headings.

We build the document with python-docx at request time (no Node dependency in
the backend image).
"""

from __future__ import annotations

import base64
import io
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.models.paper_schemas import PaperFigure, PaperResult, PaperTable, float_placements

_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"]
_FONT = "Times New Roman"
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _set_columns(section, num: int, space_twips: int = 360) -> None:
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def _add_runs_with_bold(paragraph, text: str, size: int) -> None:
    """Add text to a paragraph, rendering **bold** markdown spans as bold runs."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos : m.start()])
            r.font.size = Pt(size)
            r.font.name = _FONT
        r = paragraph.add_run(m.group(1))
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = _FONT
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(size)
        r.font.name = _FONT


def _add_figure(doc, fig: PaperFigure) -> None:
    if not fig.image_base64:
        return
    try:
        data = base64.b64decode(fig.image_base64)
    except Exception:
        return
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(io.BytesIO(data), width=Inches(3.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Fig. {fig.number}. {fig.caption}")
    cr.font.size = Pt(8)
    cr.font.name = _FONT


def _add_table(doc, tbl: PaperTable) -> None:
    if not tbl.columns:
        return
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"TABLE {tbl.number}. {tbl.caption.upper()}")
    cr.bold = True
    cr.font.size = Pt(8)
    cr.font.name = _FONT

    table = doc.add_table(rows=1, cols=len(tbl.columns))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, col in enumerate(tbl.columns):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(col))
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = _FONT
    for row in tbl.rows:
        cells = table.add_row().cells
        for j, val in enumerate(row[: len(tbl.columns)]):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = _FONT


def build_paper_docx(paper: PaperResult) -> bytes:
    """Render ``paper`` as IEEE-formatted .docx and return the bytes."""
    doc = Document()

    # Base style: Times New Roman 10pt.
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10)

    # --- Page (A4) + first section is single-column for the title block ----
    sec = doc.sections[0]
    sec.page_width = int(11906 * 635)  # A4 width in EMU (twips*635)
    sec.page_height = int(16838 * 635)
    for attr in ("top_margin", "bottom_margin"):
        setattr(sec, attr, int(0.75 * 914400))
    for attr in ("left_margin", "right_margin"):
        setattr(sec, attr, int(0.625 * 914400))
    _set_columns(sec, 1)

    # --- Title ------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(paper.title or paper.topic)
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.name = _FONT
    tr.font.color.rgb = RGBColor(0, 0, 0)

    # --- Authors ----------------------------------------------------------
    for author in paper.authors or []:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name = p.add_run(author.name)
        name.font.size = Pt(11)
        name.font.name = _FONT
        meta_bits = [b for b in (author.department, author.organization, author.city, author.email) if b]
        if meta_bits:
            p.add_run("\n")
            line = p.add_run("\n".join(meta_bits))
            line.italic = True
            line.font.size = Pt(9)
            line.font.name = _FONT

    # --- Switch to two columns for the body -------------------------------
    body = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_columns(body, 2, space_twips=360)

    # Abstract.
    if paper.abstract:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lead = p.add_run("Abstract—")
        lead.bold = True
        lead.italic = True
        lead.font.size = Pt(9)
        lead.font.name = _FONT
        rest = p.add_run(paper.abstract)
        rest.bold = True
        rest.italic = True
        rest.font.size = Pt(9)
        rest.font.name = _FONT

    # Keywords.
    if paper.keywords:
        p = doc.add_paragraph()
        lead = p.add_run("Index Terms—")
        lead.bold = True
        lead.italic = True
        lead.font.size = Pt(9)
        lead.font.name = _FONT
        kw = p.add_run(", ".join(paper.keywords))
        kw.italic = True
        kw.font.size = Pt(9)
        kw.font.name = _FONT

    # Sections, with figures/tables spread evenly through them.
    placements = float_placements(len(paper.sections), len(paper.figures), len(paper.tables))
    for i, section in enumerate(paper.sections):
        num = _ROMAN[i] if i < len(_ROMAN) else str(i + 1)
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = h.add_run(f"{num}.  {section.heading.upper()}")
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.name = _FONT
        for para_text in [p for p in section.body.split("\n") if p.strip()]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(12)
            _add_runs_with_bold(p, para_text.strip(), 10)

        for kind, idx in placements.get(i, []):
            if kind == "fig" and idx < len(paper.figures):
                _add_figure(doc, paper.figures[idx])
            elif kind == "tbl" and idx < len(paper.tables):
                _add_table(doc, paper.tables[idx])

    for kind, idx in placements.get(-1, []):
        if kind == "fig" and idx < len(paper.figures):
            _add_figure(doc, paper.figures[idx])
        elif kind == "tbl" and idx < len(paper.tables):
            _add_table(doc, paper.tables[idx])

    # References.
    if paper.references:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = h.add_run("REFERENCES")
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.name = _FONT
        for ref in paper.references:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(f"[{ref.number}] {ref.text}")
            r.font.size = Pt(8)
            r.font.name = _FONT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
